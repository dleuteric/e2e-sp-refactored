#!/usr/bin/env python3
"""
DOCX Companion Generator

Generates editable Word document template for analyst notes,
to accompany the main PDF report.

Output structure:
- Report metadata (run, target, date)
- Executive Summary (editable section)
- Key Findings (editable bullet list)
- Recommendations (editable numbered list)
- Additional Notes (free-form)
- Reference to main PDF
"""
from __future__ import annotations

import logging
from datetime import datetime
from pathlib import Path
from typing import Dict, Optional

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None

LOGGER = logging.getLogger(__name__)


class DOCXCompanion:
    """
    DOCX companion document generator.

    Creates an editable Word template with pre-formatted sections
    for analyst notes and custom content.

    Example:
        >>> companion = DOCXCompanion(
        ...     run_id="20251020T154731Z",
        ...     target_id="HGV_330",
        ...     output_path=Path("HGV_330_notes.docx")
        ... )
        >>> companion.build(kpis={'cep50_km': 0.7, 'coverage_pct': 100})
    """

    def __init__(self,
                 run_id: str,
                 target_id: str,
                 output_path: Path,
                 pdf_filename: Optional[str] = None):
        """
        Initialize DOCX companion generator.

        Args:
            run_id: Run identifier
            target_id: Target identifier
            output_path: Output DOCX path
            pdf_filename: Name of companion PDF (for reference)
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx not available. Install with: pip install python-docx"
            )

        self.run_id = run_id
        self.target_id = target_id
        self.output_path = Path(output_path)
        self.pdf_filename = pdf_filename or f"{target_id}_report.pdf"

        # Initialize document
        self.doc = Document()

        # Configure styles
        self._setup_styles()

    def _setup_styles(self):
        """Configure document-wide styles."""
        # Normal text style
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # Heading 1 (blue, large)
        if 'Heading 1' in self.doc.styles:
            h1 = self.doc.styles['Heading 1']
            h1.font.name = 'Calibri'
            h1.font.size = Pt(18)
            h1.font.bold = True
            h1.font.color.rgb = RGBColor(52, 152, 219)  # Blue

        # Heading 2 (dark gray, medium)
        if 'Heading 2' in self.doc.styles:
            h2 = self.doc.styles['Heading 2']
            h2.font.name = 'Calibri'
            h2.font.size = Pt(14)
            h2.font.bold = True
            h2.font.color.rgb = RGBColor(52, 73, 94)  # Dark gray

    def build(self, kpis: Optional[Dict[str, float]] = None) -> Path:
        """
        Build complete DOCX companion document.

        Args:
            kpis: Optional KPI values to include in metadata

        Returns:
            Path to generated DOCX file
        """
        LOGGER.info("Building DOCX companion: %s", self.output_path)

        # Ensure output directory exists
        self.output_path.parent.mkdir(parents=True, exist_ok=True)

        # Build sections
        self._add_header()
        self._add_metadata(kpis)
        self._add_section_separator()

        self._add_executive_summary_section()
        self._add_section_separator()

        self._add_key_findings_section()
        self._add_section_separator()

        self._add_recommendations_section()
        self._add_section_separator()

        self._add_additional_notes_section()
        self._add_section_separator()

        self._add_reference_section()

        # Save document
        self.doc.save(str(self.output_path))

        LOGGER.info("✓ DOCX companion saved: %s", self.output_path)
        return self.output_path

    def _add_header(self):
        """Add document title and subtitle."""
        # Main title
        title = self.doc.add_heading('ARCHITECTURE PERFORMANCE REPORT', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle (run + target)
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f'Run: {self.run_id} | Target: {self.target_id}')
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(127, 140, 141)  # Gray

        self.doc.add_paragraph()  # Spacing

    def _add_metadata(self, kpis: Optional[Dict[str, float]] = None):
        """Add metadata table (read-only info)."""
        self.doc.add_heading('Report Metadata', level=2)

        # Create 2-column table
        table = self.doc.add_table(rows=0, cols=2)
        table.style = 'Light Grid Accent 1'

        # Add rows
        def add_row(label: str, value: str):
            row = table.add_row()
            row.cells[0].text = label
            row.cells[1].text = value
            # Bold label
            row.cells[0].paragraphs[0].runs[0].font.bold = True

        add_row('Run ID', self.run_id)
        add_row('Target ID', self.target_id)
        add_row('Generated', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        add_row('Companion PDF', self.pdf_filename)

        # Add KPI summary if available
        if kpis:
            add_row('', '')  # Spacer
            for kpi_name, kpi_value in kpis.items():
                if kpi_name == 'cep50_km':
                    add_row('CEP50', f'{kpi_value:.2f} km')
                elif kpi_name == 'rmse_km':
                    add_row('RMSE', f'{kpi_value:.2f} km')
                elif kpi_name == 'coverage_pct':
                    add_row('Coverage', f'{kpi_value:.0f}%')

        self.doc.add_paragraph()  # Spacing

    def _add_section_separator(self):
        """Add visual separator between sections."""
        sep = self.doc.add_paragraph('═' * 70)
        sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sep_run = sep.runs[0]
        sep_run.font.color.rgb = RGBColor(189, 195, 199)  # Light gray
        self.doc.add_paragraph()

    def _add_executive_summary_section(self):
        """Add editable executive summary section."""
        self.doc.add_heading('EXECUTIVE SUMMARY', level=2)

        # Instruction text
        instr = self.doc.add_paragraph()
        instr_run = instr.add_run('[Edit this section with your high-level summary]')
        instr_run.italic = True
        instr_run.font.color.rgb = RGBColor(127, 140, 141)

        # Editable content area (multiple paragraphs for space)
        for _ in range(5):
            self.doc.add_paragraph()

    def _add_key_findings_section(self):
        """Add editable key findings bullet list."""
        self.doc.add_heading('KEY FINDINGS', level=2)

        # Instruction
        instr = self.doc.add_paragraph()
        instr_run = instr.add_run('[Add your key technical findings below]')
        instr_run.italic = True
        instr_run.font.color.rgb = RGBColor(127, 140, 141)

        # Pre-formatted bullet points (editable)
        for i in range(5):
            bullet = self.doc.add_paragraph(style='List Bullet')
            bullet.add_run(f'[Finding {i + 1}: describe observation...]')

    def _add_recommendations_section(self):
        """Add editable recommendations numbered list."""
        self.doc.add_heading('RECOMMENDATIONS', level=2)

        # Instruction
        instr = self.doc.add_paragraph()
        instr_run = instr.add_run('[Add actionable recommendations]')
        instr_run.italic = True
        instr_run.font.color.rgb = RGBColor(127, 140, 141)

        # Pre-formatted numbered list (editable)
        for i in range(5):
            num = self.doc.add_paragraph(style='List Number')
            num.add_run(f'[Recommendation {i + 1}: propose action...]')

    def _add_additional_notes_section(self):
        """Add free-form notes section."""
        self.doc.add_heading('ADDITIONAL NOTES', level=2)

        # Instruction
        instr = self.doc.add_paragraph()
        instr_run = instr.add_run('[Free-form notes, observations, caveats, etc.]')
        instr_run.italic = True
        instr_run.font.color.rgb = RGBColor(127, 140, 141)

        # Large editable area
        for _ in range(8):
            self.doc.add_paragraph()

    def _add_reference_section(self):
        """Add reference to main PDF report."""
        self.doc.add_heading('REFERENCE', level=2)

        ref_text = self.doc.add_paragraph()
        ref_text.add_run('Complete technical analysis: ').bold = True
        ref_text.add_run(self.pdf_filename)

        gen_text = self.doc.add_paragraph()
        gen_text.add_run('Generated: ').bold = True
        gen_text.add_run(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))


# ============================================================================
# CONVENIENCE FUNCTION
# ============================================================================
def generate_docx_companion(run_id: str,
                            target_id: str,
                            output_path: Path,
                            pdf_filename: Optional[str] = None,
                            kpis: Optional[Dict[str, float]] = None) -> Path:
    """
    Generate DOCX companion document.

    Args:
        run_id: Run identifier
        target_id: Target identifier
        output_path: Output DOCX path
        pdf_filename: Name of companion PDF
        kpis: Optional KPI values

    Returns:
        Path to generated DOCX

    Example:
        >>> docx_path = generate_docx_companion(
        ...     run_id="20251020T154731Z",
        ...     target_id="HGV_330",
        ...     output_path=Path("HGV_330_notes.docx"),
        ...     kpis={'cep50_km': 0.7, 'coverage_pct': 100}
        ... )
    """
    if not DOCX_AVAILABLE:
        LOGGER.warning("python-docx not available, skipping DOCX companion")
        return None

    companion = DOCXCompanion(
        run_id=run_id,
        target_id=target_id,
        output_path=output_path,
        pdf_filename=pdf_filename
    )

    return companion.build(kpis=kpis)


# ============================================================================
# CLI TEST
# ============================================================================
if __name__ == "__main__":
    import sys

    # Test standalone
    test_path = Path("test_companion.docx")

    try:
        docx_path = generate_docx_companion(
            run_id="20251020T154731Z",
            target_id="HGV_330",
            output_path=test_path,
            kpis={
                'cep50_km': 0.70,
                'rmse_km': 0.72,
                'coverage_pct': 100.0,
            }
        )

        print(f"✓ Test DOCX generated: {docx_path}")
        print(f"  Open with: open {docx_path}")
        sys.exit(0)

    except ImportError as e:
        print(f"✗ Missing dependency: {e}")
        print("  Install with: pip install python-docx")
        sys.exit(1)
    except Exception as e:
        print(f"✗ Error: {e}")
        import traceback

        traceback.print_exc()
        sys.exit(1)